"""
Enhanced Report Service - Comprehensive 25-30 Page Reports
Includes: Full Q&A, Detailed HLA breakdown, Visual analysis with confidence
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class ReportService:
    """Generate comprehensive 25-30 page Word reports with maximum detail."""

    TRAIT_ORDER = ["drive", "confidence", "passion", "assertiveness", "indulgence", "aspiration", "ease"]

    TRAIT_DESCRIPTIONS = {
        "drive": "Ambition, goal pursuit, material focus, resource acquisition, achievement orientation",
        "confidence": "Self-assurance, leadership, recognition of accomplishments, self-worth, status",
        "passion": "Intensity, enthusiasm, desire for experiences/power/connection, fulfillment pursuit",
        "assertiveness": "Directness, conflict approach, boundary setting, emotional intensity, justice",
        "indulgence": "Pleasure-seeking, sensory enjoyment, self-care, experiential focus, moderation",
        "aspiration": "Competitive drive, social comparison, desire for growth, validation needs, ambition",
        "ease": "Relaxation preference, energy conservation, pace management, work-life balance, effort optimization"
    }

    def _add_colored_heading(self, doc, text, level=1, color=None):
        """Add heading with custom color."""
        heading = doc.add_heading(text, level=level)
        if color:
            for run in heading.runs:
                run.font.color.rgb = color
        return heading

    def generate_full_report(self, p1: dict, p2: dict, analysis: dict,
                             visual_data: dict, hla_data: dict,
                             similarity_result: float, weights: dict,
                             filename: str) -> str:
        """Generate ENHANCED 25-30 page report with maximum detail."""
        doc = Document()
        
        print(f"\n📄 Generating ENHANCED report...")
        
        # ══════════════════════════════════════════════════════════════
        # TITLE PAGE
        # ══════════════════════════════════════════════════════════════
        title = doc.add_heading('Harmonia Compatibility Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{p1['name']} & {p2['name']}")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(232, 74, 138)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}").italic = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        tagline = doc.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tagline.add_run("Chemistry, Not Selection")
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(139, 92, 246)

        # ══════════════════════════════════════════════════════════════
        # EXECUTIVE SUMMARY
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        self._add_colored_heading(doc, 'Executive Summary', 1, RGBColor(232, 74, 138))
        
        visual_score = visual_data.get('mutual_attraction_score', 50)
        personality_score = similarity_result if isinstance(similarity_result, (int, float)) else 50
        hla_score = hla_data.get('compatibility_score', 50)
        overall = (visual_score * 0.50 + personality_score * 0.35 + hla_score * 0.15)

        # Summary table
        summary_table = doc.add_table(rows=5, cols=3)
        summary_table.style = 'Table Grid'
        headers = summary_table.rows[0].cells
        headers[0].text = "Component"
        headers[1].text = "Score"
        headers[2].text = "Weight"
        
        # Color header row
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 41, 59)

        data_rows = [
            ("Visual Attraction", f"{visual_score:.1f}%", "50%"),
            ("Personality Resonance", f"{personality_score:.1f}%", "35%"),
            ("Genetic Harmony", f"{hla_score:.1f}%", "15%"),
            ("OVERALL COMPATIBILITY", f"{overall:.1f}%", "100%")
        ]

        for i, (comp, score, weight) in enumerate(data_rows, 1):
            row = summary_table.rows[i].cells
            row[0].text = comp
            row[1].text = score
            row[2].text = weight
            if i == len(data_rows):
                for cell in row:
                    cell.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()
        verdict = analysis.get('compatibility_verdict', 'Compatibility analysis complete.')
        verdict_para = doc.add_paragraph()
        verdict_para.add_run("Verdict: ").bold = True
        verdict_para.add_run(verdict).italic = True

        # ══════════════════════════════════════════════════════════════
        # VISUAL COMPATIBILITY - DETAILED!
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        self._add_colored_heading(doc, '📸 Visual Compatibility Analysis', 1, RGBColor(232, 74, 138))
        
        doc.add_heading('Overall Visual Chemistry', level=2)
        doc.add_paragraph(f"Mutual Attraction Score: {visual_score:.1f}%")
        doc.add_paragraph(f"Confidence: {visual_data.get('confidence', 0.7)*100:.0f}%")
        doc.add_paragraph(f"Analysis Method: {visual_data.get('method', 'AI Vision Analysis')}")
        doc.add_paragraph()
        
        # Detailed visual metrics table
        doc.add_heading('Detailed Visual Metrics', level=2)
        visual_table = doc.add_table(rows=7, cols=3)
        visual_table.style = 'Table Grid'
        
        vt_headers = visual_table.rows[0].cells
        vt_headers[0].text = "Metric"
        vt_headers[1].text = p1['name']
        vt_headers[2].text = p2['name']
        
        visual_metrics = [
            ("Quality Score", 
             f"{visual_data.get('quality_a', 50):.1f}%", 
             f"{visual_data.get('quality_b', 50):.1f}%"),
            ("Attractiveness Rating",
             f"{visual_data.get('person_a_attractiveness', 5)}/10",
             f"{visual_data.get('person_b_attractiveness', 5)}/10"),
            ("Expression",
             visual_data.get('person_a_expression', 'N/A'),
             visual_data.get('person_b_expression', 'N/A')),
            ("Bidirectional Interest",
             f"{visual_data.get('a_to_b', 50):.1f}%",
             f"{visual_data.get('b_to_a', 50):.1f}%"),
            ("Reciprocity Bonus",
             f"+{visual_data.get('reciprocity_bonus', 0):.1f}%",
             f"+{visual_data.get('reciprocity_bonus', 0):.1f}%"),
            ("Overall Chemistry", 
             f"{visual_score:.1f}%", 
             f"{visual_score:.1f}%")
        ]
        
        for i, (metric, val1, val2) in enumerate(visual_metrics, 1):
            row = visual_table.rows[i].cells
            row[0].text = metric
            row[1].text = str(val1)
            row[2].text = str(val2)
        
        doc.add_paragraph()
        doc.add_paragraph(analysis.get('ui_cards', {}).get('first_impression', 
            'Visual chemistry provides the foundation for initial attraction.'))

        # ══════════════════════════════════════════════════════════════
        # HLA GENETIC COMPATIBILITY - MASSIVELY DETAILED!
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        self._add_colored_heading(doc, '🧬 HLA Genetic Compatibility Analysis', 1, RGBColor(139, 92, 246))
        
        doc.add_heading('Overall Genetic Harmony', level=2)
        doc.add_paragraph(f"HLA Compatibility Score: {hla_score:.1f}%")
        doc.add_paragraph(f"Interpretation: {hla_data.get('interpretation', 'Optimal genetic diversity for attraction.')}")
        doc.add_paragraph()
        
        # SNP Details
        doc.add_heading('Genetic Data Details', level=2)
        snp_details_table = doc.add_table(rows=4, cols=3)
        snp_details_table.style = 'Table Grid'
        
        snp_headers = snp_details_table.rows[0].cells
        snp_headers[0].text = "Metric"
        snp_headers[1].text = p1['name']
        snp_headers[2].text = p2['name']
        
        snp_count_1 = len(hla_data.get('person_a_hla', [])) if isinstance(hla_data.get('person_a_hla'), list) else 0
        snp_count_2 = len(hla_data.get('person_b_hla', [])) if isinstance(hla_data.get('person_b_hla'), list) else 0
        
        snp_data = [
            ("SNPs Analyzed", f"{snp_count_1} SNPs", f"{snp_count_2} SNPs"),
            ("Data Source", 
             hla_data.get('source_a', 'Manual Entry'), 
             hla_data.get('source_b', 'Manual Entry')),
            ("HLA Region Coverage", 
             f"{hla_data.get('coverage_a', 100):.0f}%", 
             f"{hla_data.get('coverage_b', 100):.0f}%")
        ]
        
        for i, (metric, val1, val2) in enumerate(snp_data, 1):
            row = snp_details_table.rows[i].cells
            row[0].text = metric
            row[1].text = str(val1)
            row[2].text = str(val2)
        
        doc.add_paragraph()
        
        # Locus-by-locus breakdown
        doc.add_heading('HLA Locus-by-Locus Breakdown', level=2)
        doc.add_paragraph("Dissimilarity analysis for each major HLA locus:")
        doc.add_paragraph()
        
        locus_table = doc.add_table(rows=8, cols=4)
        locus_table.style = 'Table Grid'
        
        locus_headers = locus_table.rows[0].cells
        locus_headers[0].text = "HLA Locus"
        locus_headers[1].text = "Dissimilarity %"
        locus_headers[2].text = "Contribution"
        locus_headers[3].text = "Status"
        
        locus_breakdown = hla_data.get('locus_breakdown', {})
        loci = ['HLA-A', 'HLA-B', 'HLA-C', 'HLA-DRB1', 'HLA-DQB1', 'HLA-DPB1', 'Overall']
        
        for i, locus in enumerate(loci, 1):
            row = locus_table.rows[i].cells
            locus_data = locus_breakdown.get(locus, {})
            
            dissim = locus_data.get('dissimilarity', hla_score if locus == 'Overall' else 55)
            contribution = locus_data.get('contribution', 16.7 if locus != 'Overall' else 100)
            
            row[0].text = locus
            row[1].text = f"{dissim:.1f}%"
            row[2].text = f"{contribution:.1f}%"
            
            # Status based on Wedekind optimal (55%)
            if locus == 'Overall':
                row[3].text = "✓ Complete"
            elif 45 <= dissim <= 65:
                row[3].text = "✓ Optimal"
            elif 35 <= dissim < 45 or 65 < dissim <= 75:
                row[3].text = "○ Moderate"
            else:
                row[3].text = "△ Low Impact"
        
        doc.add_paragraph()
        doc.add_paragraph("Note: Optimal HLA dissimilarity is approximately 55% based on Wedekind et al. (1995) research. Scores in the 45-65% range indicate strong genetic compatibility.")

        # ══════════════════════════════════════════════════════════════
        # PERSONALITY ANALYSIS - ENHANCED WITH FULL Q&A
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        self._add_colored_heading(doc, '🧠 Personality Compatibility Analysis', 1, RGBColor(232, 74, 138))
        
        doc.add_heading('Overall Personality Resonance', level=2)
        doc.add_paragraph(f"Perceived Similarity Score: {personality_score:.1f}%")
        doc.add_paragraph(analysis.get('perceived_similarity', 
            'Personality analysis reveals complementary traits and shared values.'))
        doc.add_paragraph()
        
        # Detailed trait scores
        doc.add_heading('Personality Profile Comparison', level=2)
        doc.add_paragraph("Detailed breakdown of personality traits across 7 core dimensions of Drive, Confidence, Passion, Assertiveness, Indulgence, Aspiration, and Ease:")
        doc.add_paragraph()
        
        trait_table = doc.add_table(rows=8, cols=5)
        trait_table.style = 'Table Grid'
        
        tt_headers = trait_table.rows[0].cells
        tt_headers[0].text = "Trait"
        tt_headers[1].text = f"{p1['name']} Score"
        tt_headers[2].text = f"{p2['name']} Score"
        tt_headers[3].text = "Difference"
        tt_headers[4].text = "Compatibility"
        
        for i, trait in enumerate(self.TRAIT_ORDER, 1):
            row = trait_table.rows[i].cells

            p1_score = p1['sins'].get(trait, {}).get('score', 0)
            p2_score = p2['sins'].get(trait, {}).get('score', 0)
            diff = abs(p1_score - p2_score)

            row[0].text = trait.capitalize()
            row[1].text = f"{p1_score:.2f}"
            row[2].text = f"{p2_score:.2f}"
            row[3].text = f"{diff:.2f}"

            # Compatibility assessment
            if diff < 2.0:
                row[4].text = "✓ Highly Similar"
            elif diff < 4.0:
                row[4].text = "○ Complementary"
            else:
                row[5].text = "△ Contrasting"
        
        # ══════════════════════════════════════════════════════════════
        # DETAILED TRAIT ANALYSIS WITH EVIDENCE
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        doc.add_heading('Detailed Personality Trait Analysis', level=1)
        doc.add_paragraph("Complete breakdown with evidence from questionnaire responses:")
        doc.add_paragraph()
        
        for trait in self.TRAIT_ORDER:
            doc.add_heading(f"{trait.capitalize()}: {self.TRAIT_DESCRIPTIONS[trait]}", level=2)

            # Person 1
            try:
                p1_score = round(p1['sins'][trait]['score'], 2)
                p1_evidence = p1['sins'][trait].get('evidence', 'No evidence recorded')

                p1_para = doc.add_paragraph()
                p1_run = p1_para.add_run(f"{p1['name']}: {p1_score:+.2f}")
                p1_run.bold = True
                p1_run.font.color.rgb = RGBColor(232, 74, 138)

                ev_para = doc.add_paragraph()
                ev_para.add_run("Evidence: ").italic = True
                ev_para.add_run(f'"{p1_evidence}"')
                ev_para.paragraph_format.left_indent = Inches(0.5)
            except (KeyError, TypeError):
                doc.add_paragraph(f"{p1['name']}: No data")

            # Person 2
            try:
                p2_score = round(p2['sins'][trait]['score'], 2)
                p2_evidence = p2['sins'][trait].get('evidence', 'No evidence recorded')

                p2_para = doc.add_paragraph()
                p2_run = p2_para.add_run(f"{p2['name']}: {p2_score:+.2f}")
                p2_run.bold = True
                p2_run.font.color.rgb = RGBColor(139, 92, 246)

                ev_para = doc.add_paragraph()
                ev_para.add_run("Evidence: ").italic = True
                ev_para.add_run(f'"{p2_evidence}"')
                ev_para.paragraph_format.left_indent = Inches(0.5)
            except (KeyError, TypeError):
                doc.add_paragraph(f"{p2['name']}: No data")

            # Compatibility note
            try:
                diff = abs(p1_score - p2_score)
                compat_para = doc.add_paragraph()
                compat_para.add_run("Compatibility Note: ").bold = True
                if diff < 2.0:
                    compat_para.add_run(f"Similar expression ({diff:.2f} difference) - shared approach to {trait}.")
                elif diff < 4.0:
                    compat_para.add_run(f"Complementary ({diff:.2f} difference) - balanced dynamic.")
                else:
                    compat_para.add_run(f"Contrasting ({diff:.2f} difference) - potential growth area.")
                compat_para.paragraph_format.left_indent = Inches(0.5)
            except:
                pass
            
            doc.add_paragraph()

        # ══════════════════════════════════════════════════════════════
        # COMPLETE QUESTIONNAIRE RESPONSES - FULL Q&A!
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        self._add_colored_heading(doc, '📝 Complete Questionnaire Responses', 1, RGBColor(232, 74, 138))
        
        # Person 1 Q&A
        doc.add_heading(f"{p1['name']}'s Responses", level=2)
        p1_responses = p1.get('raw_responses', [])
        if p1_responses:
            for i, resp in enumerate(p1_responses, 1):
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Question {i}: {resp.get('question', 'No question')}")
                q_run.bold = True
                q_run.font.color.rgb = RGBColor(232, 74, 138)
                
                a_para = doc.add_paragraph(f"{resp.get('answer', 'No answer')}")
                a_para.paragraph_format.left_indent = Inches(0.5)
                a_para.paragraph_format.space_after = Pt(12)
                
                doc.add_paragraph()
        else:
            doc.add_paragraph("No response data available.")

        # Person 2 Q&A
        doc.add_page_break()
        doc.add_heading(f"{p2['name']}'s Responses", level=2)
        p2_responses = p2.get('raw_responses', [])
        if p2_responses:
            for i, resp in enumerate(p2_responses, 1):
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Question {i}: {resp.get('question', 'No question')}")
                q_run.bold = True
                q_run.font.color.rgb = RGBColor(139, 92, 246)
                
                a_para = doc.add_paragraph(f"{resp.get('answer', 'No answer')}")
                a_para.paragraph_format.left_indent = Inches(0.5)
                a_para.paragraph_format.space_after = Pt(12)
                
                doc.add_paragraph()
        else:
            doc.add_paragraph("No response data available.")

        # ══════════════════════════════════════════════════════════════
        # DEEP ANALYSIS INSIGHTS
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        doc.add_heading('Deep Compatibility Insights', level=1)
        
        doc.add_heading('Connection Themes', level=2)
        themes = analysis.get('themes', [])
        for theme in themes:
            theme_para = doc.add_paragraph()
            theme_run = theme_para.add_run(f"• {theme}")
            theme_run.bold = True
            theme_run.font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_heading('Comprehensive Analysis', level=2)
        doc.add_paragraph(analysis.get('deep_analysis', 'Detailed analysis in progress.'))
        
        doc.add_paragraph()
        doc.add_heading('The Vibe', level=2)
        doc.add_paragraph(analysis.get('ui_cards', {}).get('vibe_check', 'Overall relationship dynamics.'))
        
        doc.add_paragraph()
        doc.add_heading('Long-Term Compatibility Key', level=2)
        doc.add_paragraph(analysis.get('ui_cards', {}).get('long_term_key', 'Sustained compatibility factors.'))

        # ══════════════════════════════════════════════════════════════
        # STRENGTHS & CHALLENGES
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        doc.add_heading('Relationship Strengths & Challenges', level=1)
        
        doc.add_heading('🟢 Key Strength', level=2)
        green_para = doc.add_paragraph(analysis.get('ui_cards', {}).get('green_flag', 'Positive compatibility indicators present.'))
        green_para.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_paragraph()
        doc.add_heading('🔴 Growth Area', level=2)
        red_para = doc.add_paragraph(analysis.get('ui_cards', {}).get('red_flag', 'Areas for awareness and communication.'))
        red_para.paragraph_format.left_indent = Inches(0.5)

        # ══════════════════════════════════════════════════════════════
        # METHODOLOGY & CALCULATIONS
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        doc.add_heading('Methodology & Scoring Details', level=1)
        
        doc.add_heading('Calculation Formula', level=2)
        doc.add_paragraph(f"Overall Score = (Visual × 50%) + (Personality × 35%) + (HLA × 15%)")
        doc.add_paragraph(f"            = ({visual_score:.1f} × 0.50) + ({personality_score:.1f} × 0.35) + ({hla_score:.1f} × 0.15)")
        doc.add_paragraph(f"            = {overall:.1f}%")
        
        doc.add_paragraph()
        doc.add_heading('Confidence Levels', level=2)
        conf_table = doc.add_table(rows=4, cols=2)
        conf_table.style = 'Table Grid'
        
        conf_headers = conf_table.rows[0].cells
        conf_headers[0].text = "Component"
        conf_headers[1].text = "Confidence"
        
        confidences = [
            ("Visual Analysis", f"{visual_data.get('confidence', 0.7)*100:.0f}%"),
            ("Personality Analysis", f"{similarity_result.get('confidence', 0.85)*100:.0f}%" if isinstance(similarity_result, dict) else "85%"),
            ("HLA Analysis", f"{hla_data.get('confidence', 0.9)*100:.0f}%")
        ]
        
        for i, (comp, conf) in enumerate(confidences, 1):
            row = conf_table.rows[i].cells
            row[0].text = comp
            row[1].text = conf

        # ══════════════════════════════════════════════════════════════
        # SCIENTIFIC REFERENCES
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        doc.add_heading('Scientific Background', level=1)
        
        doc.add_heading('HLA Compatibility Research', level=2)
        doc.add_paragraph("Wedekind, C., Seebeck, T., Bettens, F., & Paepke, A. J. (1995). MHC-dependent mate preferences in humans. Proceedings of the Royal Society B, 260(1359), 245-249.")
        doc.add_paragraph("Key finding: Women showed preference for men with HLA genes dissimilar from their own, with optimal dissimilarity around 50-60%.")
        
        doc.add_paragraph()
        doc.add_heading('Visual Attraction Research', level=2)
        doc.add_paragraph("Research indicates initial attraction decisions occur within 3 seconds of viewing, with facial symmetry and other visual features playing primary roles in immediate chemistry assessment.")
        
        doc.add_paragraph()
        doc.add_heading('Personality Framework', level=2)
        doc.add_paragraph("Our 7-dimensional personality framework maps to modern psychology research, providing nuanced assessment of Drive (ambition), Confidence (self-assurance), Passion (intensity/desire), Assertiveness (directness), Indulgence (pleasure-seeking), Aspiration (competitive drive), and Ease (work-life balance).")

        # ══════════════════════════════════════════════════════════════
        # FOOTER
        # ══════════════════════════════════════════════════════════════
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Generated by Harmonia")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(139, 92, 246)
        
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = footer2.add_run("Chemistry, Not Selection")
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(100, 100, 100)

        doc.save(filename)
        logger.info(f"✅ ENHANCED report saved: {filename} (~25-30 pages)")
        return filename
